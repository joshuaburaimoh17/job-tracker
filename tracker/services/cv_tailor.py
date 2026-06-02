import difflib
import json
import re
from copy import deepcopy
from datetime import date
from pathlib import Path

import anthropic
from docx import Document
from docx.text.paragraph import Paragraph as DocxParagraph

from django.conf import settings


SECTION_HEADINGS = {
    'Professional Summary',
    'Work Experience',
    'Projects',
    'Skills',
    'Education',
    'Achievements',
}


def _extract_section_text(doc, heading: str) -> str:
    collecting = False
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == heading:
            collecting = True
            continue
        if collecting:
            if text in SECTION_HEADINGS:
                break
            if text:
                lines.append(text)
    return '\n'.join(lines)


def _extract_all_text(doc) -> str:
    return '\n'.join(para.text for para in doc.paragraphs)


def _call_claude(original_summary: str, original_skills: str, original_experience: str,
                 job_description: str, role: str, company: str) -> dict:
    client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
    prompt = f"""You are a professional CV writer and ATS optimisation expert. Your job is to tailor a graduate's CV for a specific job application so it passes ATS filters and stands out to hiring managers. You MUST make meaningful, substantive changes — if the output looks similar to the input you have failed.

JOB ROLE: {role}
COMPANY: {company}

JOB DESCRIPTION:
{job_description[:3000]}

CURRENT PROFESSIONAL SUMMARY:
{original_summary}

CURRENT SKILLS:
{original_skills}

CURRENT WORK EXPERIENCE:
{original_experience}

RULES:
1. Rewrite the Professional Summary completely from scratch for this specific role and company. Reference the company name, the type of work, and the candidate's most relevant experience. Must be 3-4 sentences. Must be meaningfully different from the input — not just a few words swapped.
2. Mirror the exact language and keywords from the job description throughout the CV — ATS systems scan for exact keyword matches.
3. Reorder the Skills section to put the most relevant categories and skills first for this specific role. Within each category, put the most relevant skills at the front.
4. Rewrite the Work Experience bullet points to emphasise achievements and responsibilities that are most relevant to this role. Use action verbs. Use numbers and metrics where they already exist in the original — do not invent metrics. Mirror the language from the job description where it naturally fits.
5. Do NOT invent skills, qualifications, companies, or experience that are not already in the CV.
6. Keep the Skills format as "Category:  Skill1, Skill2, ..." with two spaces after the colon.
7. Keep Work Experience bullet points starting with a dash character.
8. In changes_made, be specific about what you changed and why — e.g. "Moved SQL to front of Technical skills because the role requires data querying. Added 'regulatory reporting' keyword to AIG bullet point 3 to match JD language."

Return ONLY valid JSON with no markdown fences, no preamble, no explanation:
{{"summary": "...", "skills": "...", "experience": "...", "changes_made": "..."}}"""

    message = client.messages.create(
        model='claude-opus-4-5',
        max_tokens=2048,
        messages=[{'role': 'user', 'content': prompt}],
    )
    raw = message.content[0].text.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw)


def _replace_para_text(para, new_text: str) -> None:
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ''


def _replace_section(content_paras: list, new_lines: list) -> None:
    if not content_paras:
        return
    new_lines = [line for line in new_lines if line.strip()]
    if not new_lines:
        return

    n_old = len(content_paras)
    n_new = len(new_lines)

    for i in range(min(n_old, n_new)):
        _replace_para_text(content_paras[i], new_lines[i])

    for para in reversed(content_paras[n_new:]):
        para._element.getparent().remove(para._element)

    if n_new > n_old:
        anchor = content_paras[n_old - 1]
        for line in new_lines[n_old:]:
            new_elem = deepcopy(anchor._element)
            anchor._element.addnext(new_elem)
            new_para = DocxParagraph(new_elem, None)
            _replace_para_text(new_para, line)
            anchor = new_para


def _save_tailored_docx(doc_path, tailored_data: dict, lead) -> str:
    doc = Document(str(doc_path))

    section_paras: dict[str, list] = {}
    current_heading = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in SECTION_HEADINGS:
            current_heading = text
            section_paras[current_heading] = []
        elif current_heading is not None:
            section_paras[current_heading].append(para)

    _replace_section(section_paras.get('Professional Summary', []), [tailored_data['summary']])
    _replace_section(section_paras.get('Skills', []), tailored_data['skills'].split('\n'))
    _replace_section(section_paras.get('Work Experience', []), tailored_data['experience'].split('\n'))

    output_dir = Path(settings.CV_APPLICATIONS_PATH)
    output_dir.mkdir(parents=True, exist_ok=True)

    company_slug = re.sub(r'[^\w]', '_', lead.company)[:30]
    role_slug = re.sub(r'[^\w]', '_', lead.role)[:30]
    filename = f"{company_slug}_{role_slug}_{date.today()}.docx"
    output_path = output_dir / filename

    doc.save(str(output_path))
    return str(output_path)


def compute_diff(original_text: str, tailored_text: str) -> list:
    orig_lines = original_text.splitlines()
    tail_lines = tailored_text.splitlines()
    matcher = difflib.SequenceMatcher(None, orig_lines, tail_lines)
    result = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for o, t in zip(orig_lines[i1:i2], tail_lines[j1:j2]):
                result.append(((o, 'same'), (t, 'same')))
        elif tag == 'replace':
            orig_chunk = orig_lines[i1:i2]
            tail_chunk = tail_lines[j1:j2]
            for k in range(max(len(orig_chunk), len(tail_chunk))):
                o = orig_chunk[k] if k < len(orig_chunk) else ''
                t = tail_chunk[k] if k < len(tail_chunk) else ''
                result.append(((o, 'removed' if o else 'same'), (t, 'added' if t else 'same')))
        elif tag == 'delete':
            for o in orig_lines[i1:i2]:
                result.append(((o, 'removed'), ('', 'same')))
        elif tag == 'insert':
            for t in tail_lines[j1:j2]:
                result.append((('', 'same'), (t, 'added')))
    return result


def tailor_cv_for_lead(lead) -> dict:
    doc_path = Path(settings.CV_BASE_PATH)
    doc = Document(str(doc_path))

    original_summary = _extract_section_text(doc, 'Professional Summary')
    original_skills = _extract_section_text(doc, 'Skills')
    original_experience = _extract_section_text(doc, 'Work Experience')
    original_text = '\n'.join([original_summary, original_skills, original_experience])

    tailored_data = _call_claude(
        original_summary=original_summary,
        original_skills=original_skills,
        original_experience=original_experience,
        job_description=lead.job_description,
        role=lead.role,
        company=lead.company,
    )

    cv_path = _save_tailored_docx(doc_path, tailored_data, lead)
    tailored_text = '\n'.join([
        tailored_data['summary'],
        tailored_data['skills'],
        tailored_data['experience'],
    ])

    return {
        'original_text': original_text,
        'tailored_text': tailored_text,
        'changes_made': tailored_data['changes_made'],
        'cv_path': cv_path,
    }
